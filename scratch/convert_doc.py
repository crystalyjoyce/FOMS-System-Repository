import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted, KeepTogether, HRFlowable

MD_PATH = r"C:\Users\Crystalyn Joyce\.gemini\antigravity\brain\631056b7-cc74-4e1e-a48f-73781fc2cb7e\PUP_Project_Management_Week_4_Deliverables_AI_LAYER.md"
DOCX_PATH = r"C:\Users\Crystalyn Joyce\.gemini\antigravity\brain\631056b7-cc74-4e1e-a48f-73781fc2cb7e\PUP_Project_Management_Week_4_Deliverables_AI_LAYER.docx"
PDF_PATH = r"C:\Users\Crystalyn Joyce\.gemini\antigravity\brain\631056b7-cc74-4e1e-a48f-73781fc2cb7e\PUP_Project_Management_Week_4_Deliverables_AI_LAYER.pdf"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_docx(md_text, output_path):
    doc = docx.Document()
    
    # Page Setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    code_lang = ""
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        
        # Parse Markdown Table
        parsed = []
        for r in table_rows:
            cols = [c.strip() for c in r.strip('|').split('|')]
            # Ignore separator lines like |---|---|
            if all(set(c) <= set('-:| ') for c in cols):
                continue
            parsed.append(cols)

        if parsed:
            num_cols = max(len(row) for row in parsed)
            tbl = doc.add_table(rows=len(parsed), cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tbl)

            for r_idx, row_data in enumerate(parsed):
                for c_idx in range(num_cols):
                    cell_val = row_data[c_idx] if c_idx < len(row_data) else ""
                    # Strip basic markdown bold
                    cell_val_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_val)
                    cell = tbl.cell(r_idx, c_idx)
                    cell.text = cell_val_clean
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9.5)
                        if r_idx == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

                    if r_idx == 0:
                        set_cell_background(cell, "1F4E78") # Deep Blue Header
                    elif r_idx % 2 == 1:
                        set_cell_background(cell, "F2F4F7") # Soft Light Grey

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        table_rows = []
        in_table = False

    def flush_code():
        nonlocal in_code, code_buf, code_lang
        if not code_buf:
            return
        code_text = '\n'.join(code_buf)
        
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F4F6F8")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(33, 37, 41)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        code_buf = []
        in_code = False

    while i < len(lines):
        line = lines[i]

        # Code block boundary
        if line.startswith("```"):
            if in_code:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code = True
                code_lang = line.replace("```", "").strip()
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table boundary
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Horizontal Rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 40)
            run.font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 120)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(41, 128, 185)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[4:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(44, 62, 80)
        elif line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[5:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(52, 73, 94)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = line.strip()[2:]
            # Process basic formatting
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)
        elif line.strip() != "":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

        i += 1

    if in_table:
        flush_table()
    if in_code:
        flush_code()

    doc.save(output_path)
    print(f"DOCX saved successfully to {output_path}")

def build_pdf(md_text, output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1F4E78'),
        spaceAfter=8,
        spaceBefore=12
    )

    h2_style = ParagraphStyle(
        'Heading2_Custom',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#2980B9'),
        spaceBefore=10,
        spaceAfter=4
    )

    h3_style = ParagraphStyle(
        'Heading3_Custom',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#2C3E50'),
        spaceBefore=8,
        spaceAfter=3
    )

    body_style = ParagraphStyle(
        'Body_Custom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#212529'),
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        'Bullet_Custom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#212529'),
        leftIndent=15,
        spaceAfter=2
    )

    code_style = ParagraphStyle(
        'Code_Custom',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=7.5,
        leading=9.5,
        textColor=colors.HexColor('#1A252C'),
        backColor=colors.HexColor('#F4F6F8'),
        borderPadding=6,
        spaceBefore=4,
        spaceAfter=6
    )

    story = []
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    in_table = False
    table_rows = []

    def clean_md_tags(txt):
        txt = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', txt)
        txt = re.sub(r'\*(.*?)\*', r'<i>\1</i>', txt)
        txt = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', txt)
        return txt

    def flush_pdf_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        
        parsed = []
        for r in table_rows:
            cols = [c.strip() for c in r.strip('|').split('|')]
            if all(set(c) <= set('-:| ') for c in cols):
                continue
            parsed.append(cols)

        if parsed:
            num_cols = max(len(row) for row in parsed)
            data = []
            for r_idx, r in enumerate(parsed):
                row_cells = []
                for c_idx in range(num_cols):
                    val = r[c_idx] if c_idx < len(r) else ""
                    val_html = clean_md_tags(val)
                    if r_idx == 0:
                        cell_style = ParagraphStyle('TH', parent=body_style, fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
                    else:
                        cell_style = ParagraphStyle('TD', parent=body_style, fontSize=7.5, leading=9.5)
                    row_cells.append(Paragraph(val_html, cell_style))
                data.append(row_cells)

            col_width = (7.0 * inch) / num_cols
            t = Table(data, colWidths=[col_width]*num_cols)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E78')),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8F9FA')]),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D3D3D3')),
            ]))
            story.append(t)
            story.append(Spacer(1, 6))

        table_rows = []
        in_table = False

    def flush_pdf_code():
        nonlocal in_code, code_buf
        if not code_buf:
            return
        code_text = '\n'.join(code_buf)
        # Escape html chars inside code
        code_text_esc = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Preformatted(code_text_esc, code_style))
        story.append(Spacer(1, 4))
        code_buf = []
        in_code = False

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                flush_pdf_code()
            else:
                if in_table:
                    flush_pdf_table()
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_pdf_table()

        if line.strip() == "---":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#B0B0B0'), spaceBefore=6, spaceAfter=6))
            i += 1
            continue

        if line.startswith("# "):
            story.append(Paragraph(clean_md_tags(line[2:].strip()), title_style))
        elif line.startswith("## "):
            story.append(Paragraph(clean_md_tags(line[3:].strip()), h2_style))
        elif line.startswith("### "):
            story.append(Paragraph(clean_md_tags(line[4:].strip()), h3_style))
        elif line.startswith("#### "):
            story.append(Paragraph(clean_md_tags(line[5:].strip()), h3_style))
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            txt = "&bull; " + clean_md_tags(line.strip()[2:])
            story.append(Paragraph(txt, bullet_style))
        elif line.strip() != "":
            story.append(Paragraph(clean_md_tags(line), body_style))

        i += 1

    if in_table:
        flush_pdf_table()
    if in_code:
        flush_pdf_code()

    doc.build(story)
    print(f"PDF saved successfully to {output_path}")

if __name__ == "__main__":
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()

    build_docx(md_content, DOCX_PATH)
    build_pdf(md_content, PDF_PATH)
